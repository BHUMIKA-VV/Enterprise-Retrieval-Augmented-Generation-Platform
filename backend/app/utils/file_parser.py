import os
from docx import Document
import fitz  # PyMuPDF


# =====================================================
# MAIN ENTRY
# =====================================================
def parse_file(file_path: str) -> str:
    """
    企业级文件解析器 v2
    """

    if not file_path:
        return ""

    if not os.path.exists(file_path):
        print(f"[file_parser] file not found: {file_path}")
        return ""

    ext = os.path.splitext(file_path)[1].lower()

    try:

        # =========================
        # TXT / MD
        # =========================
        if ext in [".txt", ".md"]:
            return _parse_text(file_path)

        # =========================
        # DOCX
        # =========================
        if ext == ".docx":
            return _parse_docx(file_path)

        # =========================
        # PDF
        # =========================
        if ext == ".pdf":
            return _parse_pdf(file_path)

        # =========================
        # UNKNOWN TYPE
        # =========================
        print(f"[file_parser] unsupported format: {ext}")
        return ""

    except Exception as e:
        print(f"[file_parser ERROR] {file_path}: {e}")
        return ""


# =====================================================
# TXT / MD
# =====================================================
def _parse_text(file_path: str) -> str:

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        return _clean_text(text)

    except Exception as e:
        print(f"[TXT ERROR] {e}")
        return ""


# =====================================================
# DOCX (TEXT + TABLES)
# =====================================================
def _parse_docx(file_path: str) -> str:

    try:
        doc = Document(file_path)

        parts = []

        # ---------------------
        # paragraphs
        # ---------------------
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        # ---------------------
        # tables (重要升级)
        # ---------------------
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return _clean_text("\n".join(parts))

    except Exception as e:
        print(f"[DOCX ERROR] {e}")
        return ""


# =====================================================
# PDF (PyMuPDF)
# =====================================================
def _parse_pdf(file_path: str) -> str:

    try:
        doc = fitz.open(file_path)

        text = []

        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text.append(page_text)

        return _clean_text("\n".join(text))

    except Exception as e:
        print(f"[PDF ERROR] {e}")
        return ""


# =====================================================
# TEXT CLEANER
# =====================================================
def _clean_text(text: str) -> str:

    if not text:
        return ""

    # 去掉多余空行
    lines = [line.strip() for line in text.split("\n")]

    # 过滤空行
    lines = [l for l in lines if l]

    return "\n".join(lines)